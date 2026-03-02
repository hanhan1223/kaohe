"""数据管理 API 路由"""
from fastapi import APIRouter, Depends, HTTPException, status, BackgroundTasks, UploadFile, File, Form
from sqlalchemy.orm import Session
from sqlalchemy import func, desc
from typing import List, Optional
from pydantic import BaseModel
import tempfile
import os
from pathlib import Path

from app.db.database import get_db
from app.db.models import News, Document, DocumentChunk
from app.services.crawler.odaily_crawler import OdailyCrawler, clean_news_data
from app.services.rag.rag_service import RAGService

router = APIRouter(prefix="/api/v1/data", tags=["Data Management"])


# ==================== 请求/响应模型 ====================

class CrawlNewsRequest(BaseModel):
    """爬取快讯请求"""
    pages: int = 5
    
    class Config:
        json_schema_extra = {
            "example": {
                "pages": 5
            }
        }


class AddNewsRequest(BaseModel):
    """手动添加快讯请求"""
    title: str
    content: str
    source_url: Optional[str] = None
    published_at: Optional[str] = None
    
    class Config:
        json_schema_extra = {
            "example": {
                "title": "比特币突破 10 万美元",
                "content": "据最新消息，比特币价格今日突破 10 万美元大关，创历史新高...",
                "source_url": "https://example.com/news/123",
                "published_at": "2024-01-01T10:00:00"
            }
        }


class AddNewsResponse(BaseModel):
    """添加快讯响应"""
    news_id: int
    message: str


class CrawlNewsResponse(BaseModel):
    """爬取快讯响应"""
    total_crawled: int
    total_saved: int
    message: str


class AddDocumentRequest(BaseModel):
    """添加文档请求"""
    title: str
    content: str
    source: str
    doc_type: str = "article"
    metadata: Optional[dict] = None
    
    class Config:
        json_schema_extra = {
            "example": {
                "title": "比特币白皮书",
                "content": "比特币是一种点对点的电子现金系统...",
                "source": "bitcoin.org",
                "doc_type": "whitepaper",
                "metadata": {"author": "Satoshi Nakamoto"}
            }
        }


class AddDocumentResponse(BaseModel):
    """添加文档响应"""
    document_id: int
    chunks_count: int
    message: str


class DocumentStatsResponse(BaseModel):
    """文档统计响应"""
    total_documents: int
    total_chunks: int
    by_type: dict
    
    class Config:
        json_schema_extra = {
            "example": {
                "total_documents": 50,
                "total_chunks": 500,
                "by_type": {
                    "whitepaper": 20,
                    "article": 25,
                    "tokenomics": 5
                }
            }
        }


class NewsStatsResponse(BaseModel):
    """快讯统计响应"""
    total_news: int
    latest_news: Optional[dict]
    
    class Config:
        json_schema_extra = {
            "example": {
                "total_news": 100,
                "latest_news": {
                    "id": 1,
                    "title": "比特币突破 10 万美元",
                    "created_at": "2024-01-01T10:00:00"
                }
            }
        }


class SearchDocumentsRequest(BaseModel):
    """搜索文档请求"""
    query: str
    k: int = 5
    doc_type: Optional[str] = None
    
    class Config:
        json_schema_extra = {
            "example": {
                "query": "比特币价格分析",
                "k": 5,
                "doc_type": "article"
            }
        }


class SearchDocumentsResponse(BaseModel):
    """搜索文档响应"""
    results: List[dict]
    total: int


# ==================== 快讯管理接口 ====================

@router.post("/news/add", response_model=AddNewsResponse)
async def add_news(
    request: AddNewsRequest,
    db: Session = Depends(get_db)
):
    """
    手动添加快讯
    
    参数：
    - title: 快讯标题（必需）
    - content: 快讯内容（必需）
    - source_url: 来源链接（可选）
    - published_at: 发布时间（可选，ISO 格式字符串）
    
    返回：
    - news_id: 快讯 ID
    - message: 添加结果消息
    """
    from datetime import datetime
    
    # 解析发布时间
    published_at = None
    if request.published_at:
        try:
            published_at = datetime.fromisoformat(request.published_at.replace('Z', '+00:00'))
        except ValueError:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="发布时间格式错误，请使用 ISO 格式（如：2024-01-01T10:00:00）"
            )
    
    # 创建快讯记录
    news = News(
        odaily_id=None,  # 手动添加的快讯没有 odaily_id
        title=request.title,
        content=request.content,
        published_at=published_at,
        source_url=request.source_url,
        raw_data={
            "source": "manual",
            "added_at": datetime.now().isoformat()
        }
    )
    
    db.add(news)
    db.commit()
    db.refresh(news)
    
    return AddNewsResponse(
        news_id=news.id,
        message=f"快讯添加成功，ID: {news.id}"
    )


@router.post("/news/crawl", response_model=CrawlNewsResponse)
async def crawl_news(
    request: CrawlNewsRequest,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db)
):
    """
    爬取 Odaily 快讯
    
    - **pages**: 爬取页数（默认 5 页）
    
    注意：爬取是异步进行的，会在后台执行
    """
    def crawl_task(pages: int):
        """后台爬取任务"""
        crawler = OdailyCrawler()
        total_crawled = 0
        total_saved = 0
        
        for page in range(1, pages + 1):
            news_list = crawler.fetch_news_list(page=page)
            total_crawled += len(news_list)
            
            for news_data in news_list:
                cleaned = clean_news_data(news_data)
                if not cleaned:
                    continue
                
                # 检查是否已存在
                existing = db.query(News).filter(
                    News.odaily_id == cleaned['odaily_id']
                ).first()
                
                if existing:
                    continue
                
                # 准备 raw_data（将 datetime 转换为字符串）
                raw_data_for_json = news_data.copy()
                if 'published_at' in raw_data_for_json and raw_data_for_json['published_at']:
                    raw_data_for_json['published_at'] = raw_data_for_json['published_at'].isoformat()
                
                # 保存
                news = News(
                    odaily_id=cleaned['odaily_id'],
                    title=cleaned['title'],
                    content=cleaned['content'],
                    published_at=cleaned.get('published_at'),
                    source_url=cleaned.get('source_url'),
                    raw_data=raw_data_for_json
                )
                db.add(news)
                total_saved += 1
            
            db.commit()
    
    # 添加到后台任务
    background_tasks.add_task(crawl_task, request.pages)
    
    return CrawlNewsResponse(
        total_crawled=0,  # 后台执行，暂时返回 0
        total_saved=0,
        message=f"爬取任务已启动，将爬取 {request.pages} 页快讯"
    )


@router.get("/news/stats", response_model=NewsStatsResponse)
async def get_news_stats(db: Session = Depends(get_db)):
    """获取快讯统计信息"""
    total = db.query(func.count(News.id)).scalar()
    
    latest = db.query(News).order_by(desc(News.created_at)).first()
    
    latest_news = None
    if latest:
        latest_news = {
            "id": latest.id,
            "title": latest.title,
            "created_at": latest.created_at.isoformat()
        }
    
    return NewsStatsResponse(
        total_news=total or 0,
        latest_news=latest_news
    )


@router.get("/news/list")
async def list_news(
    skip: int = 0,
    limit: int = 20,
    db: Session = Depends(get_db)
):
    """获取快讯列表"""
    news_list = db.query(News).order_by(
        desc(News.created_at)
    ).offset(skip).limit(limit).all()
    
    return [
        {
            "id": news.id,
            "title": news.title,
            "content": news.content[:200] + "..." if len(news.content) > 200 else news.content,
            "published_at": news.published_at,
            "created_at": news.created_at
        }
        for news in news_list
    ]


@router.post("/news/batch-add")
async def batch_add_news(
    news_list: List[AddNewsRequest],
    db: Session = Depends(get_db)
):
    """
    批量添加快讯
    
    参数：
    - news_list: 快讯列表，每个快讯包含 title、content 等字段
    
    返回：
    - total: 总数
    - success: 成功添加的数量
    - failed: 失败的数量
    - news_ids: 成功添加的快讯 ID 列表
    """
    from datetime import datetime
    
    success_count = 0
    failed_count = 0
    news_ids = []
    errors = []
    
    for idx, news_data in enumerate(news_list):
        try:
            # 解析发布时间
            published_at = None
            if news_data.published_at:
                try:
                    published_at = datetime.fromisoformat(news_data.published_at.replace('Z', '+00:00'))
                except ValueError:
                    errors.append({
                        "index": idx,
                        "title": news_data.title,
                        "error": "发布时间格式错误"
                    })
                    failed_count += 1
                    continue
            
            # 创建快讯记录
            news = News(
                odaily_id=None,
                title=news_data.title,
                content=news_data.content,
                published_at=published_at,
                source_url=news_data.source_url if news_data.source_url else None,
                raw_data={
                    "source": "manual_batch",
                    "added_at": datetime.now().isoformat(),
                    "has_source_url": bool(news_data.source_url),
                    "has_published_at": bool(published_at)
                }
            )
            
            db.add(news)
            db.flush()  # 获取 ID 但不提交
            
            news_ids.append(news.id)
            success_count += 1
        
        except Exception as e:
            errors.append({
                "index": idx,
                "title": news_data.title,
                "error": str(e)
            })
            failed_count += 1
    
    # 提交所有成功的记录
    if success_count > 0:
        db.commit()
    
    return {
        "total": len(news_list),
        "success": success_count,
        "failed": failed_count,
        "news_ids": news_ids,
        "errors": errors if errors else None,
        "message": f"批量添加完成：成功 {success_count} 条，失败 {failed_count} 条"
    }


@router.delete("/news/{news_id}")
async def delete_news(
    news_id: int,
    db: Session = Depends(get_db)
):
    """
    删除快讯
    
    参数：
    - news_id: 快讯 ID
    
    返回：
    - message: 删除结果消息
    """
    news = db.query(News).filter(News.id == news_id).first()
    
    if not news:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"快讯不存在: {news_id}"
        )
    
    db.delete(news)
    db.commit()
    
    return {
        "message": f"快讯删除成功: {news_id}",
        "news_id": news_id
    }


# ==================== 文档管理接口 ====================

@router.post("/documents/add", response_model=AddDocumentResponse)
async def add_document(
    request: AddDocumentRequest,
    background_tasks: BackgroundTasks
):
    """
    添加文档到 RAG 知识库
    
    - **title**: 文档标题
    - **content**: 文档内容
    - **source**: 文档来源
    - **doc_type**: 文档类型（whitepaper/article/tokenomics）
    - **metadata**: 元数据（可选）
    
    注意：向量化是异步进行的，会在后台执行
    """
    def add_task():
        """后台添加任务"""
        rag_service = RAGService()
        rag_service.add_document(
            title=request.title,
            content=request.content,
            source=request.source,
            doc_type=request.doc_type,
            metadata=request.metadata or {}
        )
    
    # 添加到后台任务
    background_tasks.add_task(add_task)
    
    return AddDocumentResponse(
        document_id=0,  # 后台执行，暂时返回 0
        chunks_count=0,
        message="文档添加任务已启动，正在向量化..."
    )


@router.get("/documents/stats", response_model=DocumentStatsResponse)
async def get_document_stats(db: Session = Depends(get_db)):
    """获取文档统计信息"""
    total_docs = db.query(func.count(Document.id)).scalar()
    total_chunks = db.query(func.count(DocumentChunk.id)).scalar()
    
    # 按类型统计
    by_type = {}
    type_stats = db.query(
        Document.doc_type,
        func.count(Document.id)
    ).group_by(Document.doc_type).all()
    
    for doc_type, count in type_stats:
        by_type[doc_type] = count
    
    return DocumentStatsResponse(
        total_documents=total_docs or 0,
        total_chunks=total_chunks or 0,
        by_type=by_type
    )


@router.get("/documents/list")
async def list_documents(
    skip: int = 0,
    limit: int = 20,
    doc_type: Optional[str] = None,
    db: Session = Depends(get_db)
):
    """获取文档列表"""
    query = db.query(Document)
    
    if doc_type:
        query = query.filter(Document.doc_type == doc_type)
    
    docs = query.order_by(desc(Document.created_at)).offset(skip).limit(limit).all()
    
    return [
        {
            "id": doc.id,
            "title": doc.title,
            "source": doc.source,
            "doc_type": doc.doc_type,
            "content_preview": doc.content[:200] + "..." if len(doc.content) > 200 else doc.content,
            "created_at": doc.created_at
        }
        for doc in docs
    ]


@router.post("/documents/search", response_model=SearchDocumentsResponse)
async def search_documents(request: SearchDocumentsRequest):
    """
    搜索文档（语义搜索）
    
    - **query**: 搜索查询
    - **k**: 返回结果数量
    - **doc_type**: 文档类型过滤（可选）
    """
    try:
        rag_service = RAGService()
        results = rag_service.search(
            query=request.query,
            k=request.k,
            doc_type=request.doc_type
        )
        
        return SearchDocumentsResponse(
            results=results,
            total=len(results)
        )
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"搜索失败: {str(e)}"
        )


@router.delete("/documents/{document_id}")
async def delete_document(
    document_id: int,
    db: Session = Depends(get_db)
):
    """删除文档"""
    doc = db.query(Document).filter(Document.id == document_id).first()
    
    if not doc:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"文档 {document_id} 不存在"
        )
    
    # 删除文档（会级联删除分块）
    db.delete(doc)
    db.commit()
    
    return {"message": f"文档 {document_id} 已删除"}


# ==================== 批量操作接口 ====================

@router.post("/documents/batch-add")
async def batch_add_documents(
    documents: List[AddDocumentRequest],
    background_tasks: BackgroundTasks
):
    """
    批量添加文档
    
    接收文档列表，批量添加到知识库
    """
    def batch_add_task():
        """后台批量添加任务"""
        rag_service = RAGService()
        for doc in documents:
            try:
                rag_service.add_document(
                    title=doc.title,
                    content=doc.content,
                    source=doc.source,
                    doc_type=doc.doc_type,
                    metadata=doc.metadata or {}
                )
            except Exception as e:
                print(f"添加文档失败 {doc.title}: {e}")
    
    background_tasks.add_task(batch_add_task)
    
    return {
        "message": f"批量添加任务已启动，共 {len(documents)} 篇文档",
        "total": len(documents)
    }


# ==================== 文件上传接口 ====================

@router.post("/documents/upload")
async def upload_document_file(
    file: UploadFile = File(...),
    title: Optional[str] = Form(None),
    source: Optional[str] = Form(None),
    doc_type: str = Form("article"),
    db: Session = Depends(get_db)
):
    """
    上传文档文件并自动切分存入向量数据库
    
    支持的文件格式：
    - .txt - 纯文本文件
    - .md - Markdown 文件
    - .pdf - PDF 文档
    - .docx - Word 文档
    
    参数：
    - file: 上传的文件
    - title: 文档标题（可选，默认使用文件名）
    - source: 文档来源（可选，默认使用 "upload"）
    - doc_type: 文档类型（可选，默认 "article"）
    
    返回：
    - document_id: 文档 ID
    - chunks_count: 切分的块数
    - message: 处理消息
    """
    
    # 验证文件格式
    allowed_extensions = {'.txt', '.md', '.pdf', '.docx'}
    file_ext = Path(file.filename).suffix.lower()
    
    if file_ext not in allowed_extensions:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"不支持的文件格式: {file_ext}。支持的格式: {', '.join(allowed_extensions)}"
        )
    
    # 使用文件名作为默认标题
    if not title:
        title = Path(file.filename).stem
    
    # 使用 "upload" 作为默认来源
    if not source:
        source = "upload"
    
    try:
        # 读取文件内容
        content_bytes = await file.read()
        
        # 根据文件类型解析内容
        if file_ext in {'.txt', '.md'}:
            # 文本文件直接解码
            content = content_bytes.decode('utf-8')
        
        elif file_ext == '.pdf':
            # PDF 文件需要 PyPDF2 或 pdfplumber
            try:
                import PyPDF2
                import io
                
                pdf_file = io.BytesIO(content_bytes)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                content = ""
                for page in pdf_reader.pages:
                    content += page.extract_text() + "\n"
                
                if not content.strip():
                    raise HTTPException(
                        status_code=status.HTTP_400_BAD_REQUEST,
                        detail="PDF 文件内容为空或无法提取文本"
                    )
            
            except ImportError:
                raise HTTPException(
                    status_code=status.HTTP_501_NOT_IMPLEMENTED,
                    detail="PDF 解析功能未安装。请安装: pip install PyPDF2"
                )
        
        elif file_ext == '.docx':
            # Word 文件需要 python-docx
            try:
                import docx
                import io
                
                doc_file = io.BytesIO(content_bytes)
                doc = docx.Document(doc_file)
                
                content = ""
                for paragraph in doc.paragraphs:
                    content += paragraph.text + "\n"
                
                if not content.strip():
                    raise HTTPException(
                        status_code=status.HTTP_400_BAD_REQUEST,
                        detail="Word 文件内容为空"
                    )
            
            except ImportError:
                raise HTTPException(
                    status_code=status.HTTP_501_NOT_IMPLEMENTED,
                    detail="Word 解析功能未安装。请安装: pip install python-docx"
                )
        
        # 验证内容不为空
        if not content.strip():
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="文件内容为空"
            )
        
        # 使用 RAG 服务添加文档
        rag_service = RAGService()
        document_id, chunks_count = rag_service.add_document(
            title=title,
            content=content,
            source=source,
            doc_type=doc_type,
            metadata={
                "filename": file.filename,
                "file_type": file_ext,
                "file_size": len(content_bytes)
            }
        )
        
        return {
            "document_id": document_id,
            "chunks_count": chunks_count,
            "filename": file.filename,
            "title": title,
            "content_length": len(content),
            "message": f"文件上传成功，已切分为 {chunks_count} 个块并存入向量数据库"
        }
    
    except UnicodeDecodeError:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="文件编码错误，请确保文本文件使用 UTF-8 编码"
        )
    
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"文件处理失败: {str(e)}"
        )


@router.post("/documents/upload-batch")
async def upload_multiple_documents(
    files: List[UploadFile] = File(...),
    doc_type: str = Form("article"),
    db: Session = Depends(get_db)
):
    """
    批量上传文档文件
    
    支持同时上传多个文件，每个文件独立处理
    
    参数：
    - files: 上传的文件列表
    - doc_type: 文档类型（应用于所有文件）
    
    返回：
    - total: 总文件数
    - success: 成功处理的文件数
    - failed: 失败的文件数
    - results: 每个文件的处理结果
    """
    
    results = []
    success_count = 0
    failed_count = 0
    
    for file in files:
        try:
            # 验证文件格式
            allowed_extensions = {'.txt', '.md', '.pdf', '.docx'}
            file_ext = Path(file.filename).suffix.lower()
            
            if file_ext not in allowed_extensions:
                results.append({
                    "filename": file.filename,
                    "status": "failed",
                    "error": f"不支持的文件格式: {file_ext}"
                })
                failed_count += 1
                continue
            
            # 读取和解析文件
            content_bytes = await file.read()
            
            if file_ext in {'.txt', '.md'}:
                content = content_bytes.decode('utf-8')
            
            elif file_ext == '.pdf':
                try:
                    import PyPDF2
                    import io
                    
                    pdf_file = io.BytesIO(content_bytes)
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    content = ""
                    for page in pdf_reader.pages:
                        content += page.extract_text() + "\n"
                except ImportError:
                    results.append({
                        "filename": file.filename,
                        "status": "failed",
                        "error": "PDF 解析功能未安装"
                    })
                    failed_count += 1
                    continue
            
            elif file_ext == '.docx':
                try:
                    import docx
                    import io
                    
                    doc_file = io.BytesIO(content_bytes)
                    doc = docx.Document(doc_file)
                    content = ""
                    for paragraph in doc.paragraphs:
                        content += paragraph.text + "\n"
                except ImportError:
                    results.append({
                        "filename": file.filename,
                        "status": "failed",
                        "error": "Word 解析功能未安装"
                    })
                    failed_count += 1
                    continue
            
            if not content.strip():
                results.append({
                    "filename": file.filename,
                    "status": "failed",
                    "error": "文件内容为空"
                })
                failed_count += 1
                continue
            
            # 添加文档
            title = Path(file.filename).stem
            rag_service = RAGService()
            document_id, chunks_count = rag_service.add_document(
                title=title,
                content=content,
                source=file.filename,  # 使用文件名作为 source，避免重复
                doc_type=doc_type,
                metadata={
                    "filename": file.filename,
                    "file_type": file_ext,
                    "file_size": len(content_bytes)
                }
            )
            
            results.append({
                "filename": file.filename,
                "status": "success",
                "document_id": document_id,
                "chunks_count": chunks_count
            })
            success_count += 1
        
        except Exception as e:
            results.append({
                "filename": file.filename,
                "status": "failed",
                "error": str(e)
            })
            failed_count += 1
    
    return {
        "total": len(files),
        "success": success_count,
        "failed": failed_count,
        "results": results,
        "message": f"批量上传完成：成功 {success_count} 个，失败 {failed_count} 个"
    }

